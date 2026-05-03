"""Generate 论文初稿_v5.docx from v5.md using python-docx.

One-shot script. Run from project root.
"""

import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os


def markdown_to_docx(md_path, docx_path):
    """Convert the v5 Markdown paper to a formatted .docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── Style helpers ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(0.74)

    i = 0
    in_table = False
    in_code_block = False
    in_quote = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Heading detection ──
        heading_match = re.match(r'^(#{1,6})\s+(.+?)(?:\s+\{#.*\})?$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown markers from heading text
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            level = min(level + 1, 6)  # shift down one level for docx
            heading = doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph('─' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ── Blank line ──
        if not line.strip():
            i += 1
            continue

        # ── Table detection ──
        table_row_match = re.match(r'^\|(.+)\|$', line)
        if table_row_match:
            # Collect all consecutive table lines
            table_lines = []
            while i < len(lines) and re.match(r'^\|(.+)\|$', lines[i].rstrip('\n')):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            # Filter out separator lines (e.g. |---|---|)
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue  # skip separator
                data_rows.append(cells)

            if data_rows:
                num_cols = len(data_rows[0])
                table = doc.add_table(rows=len(data_rows), cols=num_cols, style='Light Grid Accent 1')
                for r, row_data in enumerate(data_rows):
                    for c, cell_text in enumerate(row_data):
                        if c < num_cols:
                            cell = table.rows[r].cells[c]
                            # Bold header row
                            p = cell.paragraphs[0]
                            p.text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                            if r == 0:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                            else:
                                for run in p.runs:
                                    run.font.size = Pt(10)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # spacing after table
            continue

        # ── Blockquote ──
        if line.lstrip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].rstrip('\n').lstrip().startswith('>'):
                quote_text = lines[i].rstrip('\n').lstrip()[1:].strip()
                quote_lines.append(quote_text)
                i += 1
            quote_para = doc.add_paragraph()
            quote_para.paragraph_format.left_indent = Cm(1.0)
            quote_para.style = doc.styles['Normal']
            run = quote_para.add_run('\n'.join(quote_lines))
            run.font.size = Pt(10)
            run.italic = True
            continue

        # ── Code block ──
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip('\n').strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Cm(0.5)
                run = code_para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # ── Regular paragraph ──
        text = line.strip()
        # Handle bold markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Handle inline code
        text = re.sub(r'`([^`]+)`', r'\1', text)

        para = doc.add_paragraph(text)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(root, 'docs', '论文初稿_v5.md')
    docx_path = os.path.join(root, 'docs', '论文初稿_v5.docx')
    markdown_to_docx(md_path, docx_path)
