"""Generate Full_Draft.docx from LaTeX section files.

Converts LaTeX markup to python-docx formatted content with proper heading
hierarchy, tables, equations (as styled text), and bullet lists.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent
DRAFT_DIR = PROJECT / "paper_draft"
OUTPUT = DRAFT_DIR / "Full_Draft.docx"


def strip_latex(text: str) -> str:
    """Remove common LaTeX commands, keeping content."""
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textasciitilde\s*', '~', text)
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\ref\{([^}]*)\}', r'[\1]', text)
    text = re.sub(r'~', ' ', text)
    text = re.sub(r'\\,', ' ', text)
    text = re.sub(r'\\\\\s*', ' ', text)
    text = re.sub(r'\\\\', '', text)
    text = text.replace('\\%', '%')
    text = text.replace('\\$', '$')
    text = text.replace('\\&', '&')
    text = text.replace('\\#', '#')
    text = text.replace("``", '"').replace("''", '"')
    text = text.replace('\\rightarrow', '→')
    text = text.replace('\\leftarrow', '←')
    text = text.replace('\\leftrightarrow', '↔')
    text = text.replace('\\times', '×')
    text = text.replace('\\sim', '~')
    text = text.replace('\\Delta', 'Δ')
    text = text.replace('\\alpha', 'α')
    text = text.replace('\\tau', 'τ')
    text = text.replace('\\mathbb{R}', 'ℝ')
    text = text.replace('\\text{', '').replace('}', '')
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\ ', ' ', text)
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = re.sub(r'\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False, size=None):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(strip_latex(text))
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def parse_table_block(lines):
    """Parse a LaTeX tabular environment into header + rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('\\') or line.startswith('%') or '&' not in line:
            continue
        line = re.sub(r'\\\\.*', '', line)
        cells = [strip_latex(c.strip()) for c in line.split('&')]
        rows.append(cells)
    return rows


def add_table(doc, rows, caption=""):
    """Add a formatted table to the document."""
    if not rows:
        return
    if caption:
        cap = strip_latex(re.sub(r'\\caption\{([^}]*)\}', r'\1', caption))
        p = doc.add_paragraph()
        run = p.add_run(f"Table: {cap}")
        run.bold = True
        run.font.size = Pt(10)

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()


def process_latex_file(doc, filepath, section_level=1):
    """Parse a single .tex file and add content to the document."""
    text = filepath.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    in_table = False
    table_lines = []
    table_caption = ""
    in_itemize = False
    in_enumerate = False
    in_equation = False
    eq_lines = []
    list_counter = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip comments and empty label-only lines
        if line.startswith('%') or (line.startswith('\\label{') and line.endswith('}')):
            i += 1
            continue

        if '\\begin{abstract}' in line or '\\end{abstract}' in line:
            i += 1
            continue

        # Section headings
        sec_match = re.match(r'\\section\{(.+?)\}', line)
        if sec_match:
            doc.add_heading(strip_latex(sec_match.group(1)), level=section_level)
            i += 1
            continue

        subsec_match = re.match(r'\\subsection\{(.+?)\}', line)
        if subsec_match:
            doc.add_heading(strip_latex(subsec_match.group(1)), level=section_level + 1)
            i += 1
            continue

        para_match = re.match(r'\\paragraph\{(.+?)\}', line)
        if para_match:
            rest = line[para_match.end():].strip()
            p = doc.add_paragraph()
            run = p.add_run(strip_latex(para_match.group(1)) + ". ")
            run.bold = True
            run.font.size = Pt(11)
            if rest:
                run2 = p.add_run(strip_latex(rest))
                run2.font.size = Pt(11)
            i += 1
            continue

        # Table environment
        if '\\begin{table}' in line:
            in_table = True
            table_lines = []
            table_caption = ""
            i += 1
            continue
        if in_table:
            if '\\end{table}' in line:
                in_table = False
                rows = parse_table_block(table_lines)
                add_table(doc, rows, table_caption)
                i += 1
                continue
            if '\\caption{' in line:
                table_caption = line
            table_lines.append(line)
            i += 1
            continue

        # Equation environments
        if '\\begin{equation}' in line or '\\begin{align}' in line:
            in_equation = True
            eq_lines = []
            i += 1
            continue
        if in_equation:
            if '\\end{equation}' in line or '\\end{align}' in line:
                in_equation = False
                eq_text = ' '.join(eq_lines)
                eq_text = strip_latex(eq_text)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(eq_text)
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)
                i += 1
                continue
            eq_lines.append(line)
            i += 1
            continue

        # Itemize / enumerate
        if '\\begin{itemize}' in line:
            in_itemize = True
            i += 1
            continue
        if '\\end{itemize}' in line:
            in_itemize = False
            i += 1
            continue
        if '\\begin{enumerate}' in line:
            in_enumerate = True
            list_counter = 0
            i += 1
            continue
        if '\\end{enumerate}' in line:
            in_enumerate = False
            i += 1
            continue

        if (in_itemize or in_enumerate) and '\\item' in line:
            item_text = re.sub(r'\\item\s*', '', line).strip()
            if in_enumerate:
                list_counter += 1
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph(style='List Bullet')
            cleaned = strip_latex(item_text)
            p.add_run(cleaned)
            i += 1
            continue

        # Regular paragraph text
        if line and not line.startswith('\\'):
            para_text = line
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (not next_line or next_line.startswith('\\') or
                    next_line.startswith('%') or '\\begin' in next_line):
                    break
                para_text += ' ' + next_line
                i += 1
            cleaned = strip_latex(para_text)
            if cleaned:
                doc.add_paragraph(cleaned)
            i += 1
            continue

        i += 1


def build_document():
    doc = Document()

    # Title
    title = doc.add_heading('Distribution-Driven Child Speech Emotion Recognition', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('INTERSPEECH / ICASSP 2026 — Working Draft')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    doc.add_heading('Abstract', level=1)
    abstract_file = DRAFT_DIR / "0_Abstract.tex"
    if abstract_file.exists():
        process_latex_file(doc, abstract_file, section_level=2)

    for fname in (
        "1_Introduction.tex",
        "2_Related_Work.tex",
        "3_Methodology.tex",
        "4_Experiments_and_Results.tex",
        "5_Analysis_and_Discussion.tex",
        "6_Conclusion.tex",
    ):
        path = DRAFT_DIR / fname
        if path.exists():
            process_latex_file(doc, path, section_level=1)

    # References placeholder
    doc.add_heading('References', level=1)
    refs = [
        '[chen2022wavlm] Chen, S. et al. "WavLM: Large-Scale Self-Supervised Pre-Training for Full Stack Speech Processing." IEEE JSTSP, 2022.',
        '[fan2024children] Fan, Y. et al. "Children\'s Speech Emotion Recognition: A Comprehensive Survey." Speech Communication, 2024.',
        '[mcfee2015librosa] McFee, B. et al. "librosa: Audio and Music Signal Analysis in Python." SciPy, 2015.',
        '[wagner2023dawn] Wagner, J. et al. "Dawn of the Transformer Era in Speech Emotion Recognition." IEEE TAFFC, 2023.',
        '[pasad2021layer] Pasad, A. et al. "Layer-wise Analysis of a Self-supervised Speech Representation Model." ASRU, 2021.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.runs[0].font.size = Pt(9)

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == '__main__':
    build_document()
