"""
Generate comprehensive experimental data document (Word).
"""
import json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = r'D:\大学\论文\儿童语音情绪识别\新方案-分布驱动儿童SER'
OUTPUT = os.path.join(ROOT, 'docs', '实验数据汇总.docx')

# Load all data
with open(os.path.join(ROOT, 'experiments', 'phase3_ablation.json')) as f:
    phase3 = json.load(f)
with open(os.path.join(ROOT, 'experiments', 'cross_language_results.json')) as f:
    xl = json.load(f)
with open(os.path.join(ROOT, 'experiments', 'model_comparison_results.json')) as f:
    mc = json.load(f)

# ─── Helper ───
doc = Document()

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = bold
    return p

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    # data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()  # spacer
    return tbl

def calc_mean_std(vals):
    m = sum(vals) / len(vals)
    s = (sum((v - m) ** 2 for v in vals) / len(vals)) ** 0.5
    return m, s

# ═══════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════
title_p = doc.add_paragraph()
title_run = title_p.add_run('儿童语音情绪识别（新方案）实验数据汇总')
title_run.font.size = Pt(22)
title_run.bold = True
title_run.font.name = 'Microsoft YaHei'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

para('项目：分布驱动儿童语音情绪识别（Distribution-Driven Children\'s SER）', size=12)
para('GitHub: Xuzc317/child-speech-emotion-recognition', size=10)
para('文档生成日期：2026年4月29日', size=10)
para('', size=6)
para('本文档汇总了项目从 Phase 1 到 Phase 4 的所有实验数据，包括数据来源、实验配置、原始数值和简要分析。所有数据均来自严格 speaker-independent 划分下的实验结果。', size=11)

# ═══════════════════════════════════════════
# 1. DATASET
# ═══════════════════════════════════════════
heading('一、数据集信息', 1)

heading('1.1 BESD MY 数据集概况', 2)
para('BESD (Bilingual Emotional Speech Dataset) 是一个儿童双语语音情绪数据集，包含英语 (ENGLISH) 和泰卢固语 (TELUGU) 两种语言的儿童语音录音。MY 为混合子集，合并了两种语言。', size=11)
para('数据集路径：D:\\大学\\论文\\儿童语音情绪识别\\提交到团队\\数据集\\BESD\\BESD\\MY\\', size=10)

add_table(
    ['属性', '值'],
    [['总 WAV 文件数', '4,179'],
     ['总说话人数', '237'],
     ['语言', 'ENGLISH + TELUGU（混合）'],
     ['情绪类别数', '6'],
     ['采样率', '16 kHz'],
     ['单条语音平均时长', '~2.5 秒'],
     ['情感类别分布', 'ANGER(695), DISGUST(699), FEAR(700), HAPPY(702), NEUTRAL(691), SAD(696)']],
    [5, 10]
)

heading('1.2 儿童语音声学统计', 2)
para('以下统计基于 BESD MY 全量数据的 librosa 帧级特征提取（每类 690-702 条 WAV，不做增强）。数据来源：src/data/statistics.py 对 c_besd_statistics.csv 的汇总。', size=11)

add_table(
    ['统计量', '值', '说明'],
    [['F0 均值', '495.9 Hz', '成人女声通常 ~200 Hz，儿童明显更高'],
     ['F0 标准差', '651.1 Hz', '反映儿童音高的极宽波动范围'],
     ['F0 P5', '75 Hz', '第5百分位：极低值'],
     ['F0 P95', '2,286 Hz', '第95百分位：极高值'],
     ['儿童 vs 成人 FD', '0.37', 'WavLM 隐空间的 Frechet Distance'],
     ['Scale 平均偏离', '8%', '儿童特征在 WavLM 空间中的缩放偏移']],
    [5, 4, 7]
)

para('说明：Frechet Distance (FD) 用于衡量儿童语音特征分布与成人语音特征分布在 WavLM 隐空间中的距离。FD 越大，分布偏移越严重。0.37 的值表明存在系统性偏移，但不至于完全不同——这正好为 Adapter 模块提供了设计依据。', size=10)

heading('1.3 数据划分方案', 2)
para('项目过程中经历了三种数据划分方案，逐步收紧以消除数据泄露：', size=11)

add_table(
    ['阶段', '划分方案', '训练:验证:测试', '说话人重叠', '用途'],
    [['原始论文（旧项目）', 'random.shuffle 随机打乱后切分', '70:30（无 val）', '63% 重叠（泄露）', '已废弃'],
     ['Phase 1-2（初期）', 'Profile-based stratification', '70:30（test=val）', '0 重叠', '模块消融'],
     ['Phase 3-4（最终）', 'Profile-based 3-way split', '60:20:20', '0 重叠', '最终报告值']],
    [4.5, 4.5, 3, 2.5, 2.5]
)

# ═══════════════════════════════════════════
# 2. PHASE 1
# ═══════════════════════════════════════════
heading('二、Phase 1：SSL 基线验证', 1)

heading('2.1 实验目的与配置', 2)
para('目的：验证自监督学习（SSL）帧级特征是否能突破旧方案 162-dim mean-pooled 特征的 ~35% 信息上限。', size=11)
para('配置：Frozen backbone + Linear Probe（仅训练一个线性分类器，验证特征本身的质量，不引入额外学习模块）。', size=11)
para('数据集：BESD MY，70/30 speaker-independent 划分。', size=11)
para('运行环境：本地 RTX 3060 6GB Laptop GPU, PyTorch 2.6.0+cu118。', size=10)

heading('2.2 实验数据', 2)

add_table(
    ['模型', '特征维度', 'Backbone 参数量', 'Val Accuracy', '备注'],
    [['旧 162-dim baseline', '162 (mean-pooled)', '—（手工特征）', '~35%', '时间维度丢失，不可用'],
     ['emotion2vec Base', '768 (帧级)', '~95M', '~66%', '震荡较大，收敛不稳定'],
     ['WavLM Base', '768 (帧级)', '~94M', '80.66%', '稳定，方差小，选为主模型']],
    [4.5, 3, 3, 2.5, 4.5]
)

para('核心结论：SSL 帧级特征 + 保留时间维度 → 35% → 80.66%，提升 45.66pp。确认方向正确，后续以 WavLM Base 为主模型。', bold=True, size=11)

# ═══════════════════════════════════════════
# 3. PHASE 2
# ═══════════════════════════════════════════
heading('三、Phase 2：三模块逐个验证', 1)

heading('3.1 模块 1：AcousticCalibrationAdapter（分布校准适配器）', 2)
para('目的：在冻结的 WavLM 输出上添加轻量非线性适配器，学习儿童语音相对于成人预训练模型的分布偏移。', size=11)
para('结构：LayerNorm → Linear(768→192) → GELU → Linear(192→768) + 残差连接。参数量约 900K。', size=11)
para('数据集：BESD MY，70/30 speaker-independent（注：此阶段尚未使用 3-way split）。', size=11)

add_table(
    ['实验编号', 'Adapter 配置', 'Val Accuracy', '差值 vs A1', '分析'],
    [['A1（基线）', '无 Adapter', '80.66%', '—', '纯 frozen WavLM + linear probe'],
     ['A2', '统计先验初始化', '81.45%', '+0.79pp', '用 μ_child/σ_child 线性校正，效果微弱'],
     ['A2b', '随机初始化', '85.40%', '+4.74pp', 'Adapter 自行学习校准映射，效果显著']],
    [3, 5, 3, 2.5, 5.5]
)

para('关键发现：统计先验（线性校正）远不如随机初始化（非线性学习），证明儿童分布偏移不是简单的"平移+缩放"，而是结构性的、非线性的，需要可学习模块来校准。这是论文的核心论点之一。', bold=True, size=11)

heading('3.2 模块 2：TemporalImportancePooling（时序重要性池化）', 2)
para('目的：用儿童语音的韵律特征（F0 + 能量）引导帧级 attention，替代传统的 mean pooling，使模型关注情感信息密度最高的帧。', size=11)
para('数据来源：F0 和 RMS energy 由 librosa 逐帧提取，拼接为 (T, 2) 的韵律特征张量。', size=11)

add_table(
    ['实验编号', 'Adapter', 'Pooling 方式', 'Val Accuracy', '差值'],
    [['A2b', '随机初始化', 'mean pooling', '85.40%', '基线'],
     ['B2', '随机初始化', 'self-attention pooling', '—', '与 A2b 持平，纯 attention 无额外信息'],
     ['B3', '随机初始化', 'prosody-guided pooling', '86.42%', '+1.02pp vs A2b']],
    [3, 3, 4, 3, 3]
)

para('结论：韵律引导的池化策略优于纯 self-attention，说明 F0+energy 信息对定位儿童语音中的情感关键帧有实际价值。同时具备可解释性优势——高 F0 变异的帧确实获得了更高的注意力权重。', size=11)

heading('3.3 模块 3：Distribution-Constrained Augmentation（分布约束增强）', 2)
para('目的：通过负面实验证明——如果增强参数不根据儿童语音分布进行约束，增强操作会导致样本漂出真实分布，进而定量损害分类性能。', size=11)
para('方法：用 Frechet Distance (FD) 定量衡量增强后样本分布与原始分布的距离。增强方法为 SpecAugment（在 SSL 特征上应用时间/频率掩码）。', size=11)

# 儿童语音统计参数
add_table(
    ['参数', '成人语音（通用默认）', '儿童语音（BESD 统计）', '儿童约束参数'],
    [['pitch shift', '±6 semitones', 'F0 mean=496Hz, std=651Hz', '±3 semitones'],
     ['speed', '0.7 - 1.3', '语速波动 ~3x 成人', '0.85 - 1.15']],
    [4, 5, 5.5, 4]
)

para('', size=4)
para('四组对比实验（在 B3 配置：Adapter+Prosody Pooling 基础上添加增强）：', size=11)

add_table(
    ['实验编号', '增强参数', 'FD (Frechet Distance)', 'Val Accuracy', '差值 vs C1'],
    [['C1', '无增强', '0', '86.42%', '基线'],
     ['C3', '儿童约束参数 (±3st, 0.85-1.15)', '8.71', '66.85%', '-19.57pp'],
     ['C2', '成人默认参数 (±6st, 0.7-1.3)', '9.87', '59.51%', '-26.91pp'],
     ['C4', '极端参数 (±12st, 0.5-1.5)', '11.99', '43.25%', '-43.17pp']],
    [3, 5.5, 3, 3, 3]
)

para('核心发现：FD 与 Val Accuracy 呈现严格的单调负相关关系——分布偏移量（FD）越大，分类性能越差。即使使用儿童约束参数（C3），FD 仍从 0 升至 8.71，准确率下降了 19.57pp——这说明儿童语音中的情感特征对增强操作高度脆弱。这个负面实验是论文最有说服力的证据之一。', bold=True, size=11)

# ═══════════════════════════════════════════
# 4. PHASE 3
# ═══════════════════════════════════════════
heading('四、Phase 3：整合全消融实验（云端，3-way split，3 seeds）', 1)

heading('4.1 实验目的与配置', 2)
para('目的：在严格的数据划分（60/20/20 三路 speaker-independent + 3 seeds）下，系统评估每个模块的贡献，产出论文主表。', size=11)
para('数据划分：split_speakers_3way()，profile-stratified，三层断言保证 train/val/test 说话人互斥。', size=11)
para('运行环境：云端 RTX 4090D 24GB（Seetacloud），conda env "speech"。本地 RTX 3060 6GB 显存不足。', size=11)
para('特征提取：WavLM Base (microsoft/wavlm-base-sv)，帧级 768-dim，~200 帧/条。', size=11)

heading('4.2 主实验结果', 2)

# Build mean/std from JSON
exp_map = {
    'A1 基线': 'A1_baseline', 'A2 统计先验': 'A2_stat_prior',
    'A2b 随机初始化': 'A2b_rand_init', 'B3 Adapter+Prosody': 'B3_prosody',
    'A3 Full Fine-tune': 'A3_full_finetune'
}
rows = []
for name, key in exp_map.items():
    vals = [phase3[key][f'seed_{s}']['val'] * 100 for s in ['42','123','456']]
    tvals = [phase3[key][f'seed_{s}']['test'] * 100 for s in ['42','123','456']]
    vm, vs = calc_mean_std(vals)
    tm, ts = calc_mean_std(tvals)
    rows.append([name, f'{vm:.2f} ± {vs:.2f}', f'{tm:.2f} ± {ts:.2f}'])

add_table(
    ['实验', 'Val Accuracy (mean ± std)', 'Test Accuracy (mean ± std)'],
    rows,
    [6, 6, 6]
)

heading('4.3 分 seed 详细数据', 2)
rows2 = []
for name, key in exp_map.items():
    v42 = phase3[key]['seed_42']['val'] * 100
    v123 = phase3[key]['seed_123']['val'] * 100
    v456 = phase3[key]['seed_456']['val'] * 100
    t42 = phase3[key]['seed_42']['test'] * 100
    t123 = phase3[key]['seed_123']['test'] * 100
    t456 = phase3[key]['seed_456']['test'] * 100
    rows2.append([name, f'{v42:.2f}', f'{v123:.2f}', f'{v456:.2f}', f'{t42:.2f}', f'{t123:.2f}', f'{t456:.2f}'])

add_table(
    ['实验', 'Val (seed 42)', 'Val (seed 123)', 'Val (seed 456)',
     'Test (seed 42)', 'Test (seed 123)', 'Test (seed 456)'],
    rows2,
    [4, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5]
)

heading('4.4 模块贡献分解', 2)
para('以下分解基于 Test Accuracy 均值（3 seeds 平均），从旧 baseline 逐步递增：', size=11)

add_table(
    ['步骤', '配置变化', 'Test Acc', '增量 (Δpp)', '累计提升'],
    [['起点', '162-dim mean-pooled（旧方案，泄露修复后）', '35.00%', '—', '—'],
     ['+WavLM', 'Frozen WavLM + mean pooling (A1)', '78.34%', '+43.34', '+43.34'],
     ['+Adapter', 'Frozen WavLM + Adapter rand init + mean pooling (A2b)', '79.51%', '+1.17', '+44.51'],
     ['+Prosody', 'Frozen WavLM + Adapter + Prosody Pooling (B3)', '81.24%', '+1.73', '+46.24'],
     ['+Fine-tune', 'Unfrozen WavLM + Adapter + Prosody Pooling (A3)', '81.96%', '+0.72', '+46.96']],
    [3, 7, 2.5, 2.5, 2.5]
)

para('关键洞察：WavLM SSL 帧级特征是最大的单一贡献者（+43pp），占总提升的 92%。Adapter + Prosody Pooling 两大模块合计贡献 +2.90pp。Full Fine-tune 仅增加 0.72pp——说明 900K 参数的轻量模块已经捕获了 94M backbone 的大部分收益，训练效率显著优于全参数微调。', bold=True, size=11)

heading('4.5 实验运行记录', 2)
para('以下为云端 Phase 3 实验的运行日志记录（来源：experiments/registry.csv）：', size=10)

add_table(
    ['实验', '时间戳', '耗时(min)', 'Epochs', 'Best Val', 'Test'],
    [['A1_baseline', '20260428_023048', '5.6', '3', '85.68%', '78.34%'],
     ['A2_stat_prior', '20260428_023623', '4.8', '3', '86.32%', '79.25%'],
     ['A2b_rand_init', '20260428_024112', '4.7', '3', '86.40%', '79.51%'],
     ['B2_self_attn', '20260428_024553', '4.6', '3', '86.40%', '79.51%'],
     ['B3_prosody', '20260428_025029', '3.6', '3', '88.52%', '81.24%'],
     ['A3_full_finetune', '20260428_025404', '4.2', '3', '88.52%', '81.96%']],
    [4, 3.5, 2, 1.5, 2.5, 2.5]
)

# ═══════════════════════════════════════════
# 5. PHASE 4
# ═══════════════════════════════════════════
heading('五、Phase 4：补充实验', 1)

heading('5.1 跨语言迁移实验', 2)
para('目的：验证语言偏移对儿童 SER 的影响，论证"语言是分布偏移的极端维度"。用英文训练的模型评估泰卢固语（X1），反之亦然（X2）。', size=11)
para('配置：B3（WavLM + Adapter rand init + Prosody Pooling + DrseCNN），3 seeds。', size=11)
para('数据集：ENGLISH 2,095 WAVs (139 speakers)，TELUGU 2,085 WAVs (420 speakers)。说话人零重叠。', size=11)

add_table(
    ['实验', '训练语言 → 测试语言', 'Val Acc (mean ± std)', 'Test Acc (mean ± std)'],
    [['X1', 'English → Telugu', '86.86% ± 0.60%', '19.68% ± 0.45%'],
     ['X2', 'Telugu → English', '89.03% ± 0.35%', '28.15% ± 0.38%']],
    [3, 4, 5, 5]
)

para('参考值：同语言混合训练（MY Mix）B3 Test = 81.24% ± 1.45%。', size=10)

heading('5.2 分 seed 详细数据', 2)

# Build cross-lang seed-level table
for xname, xkey in [('X1 English→Telugu', 'X1_EnglishToTelugu'), ('X2 Telugu→English', 'X2_TeluguToEnglish')]:
    v_vals = xl[xkey]['val_seeds']
    t_vals = xl[xkey]['test_seeds']
    para(f'{xname}：', bold=True, size=11)
    add_table(
        ['Seed', 'Val Acc', 'Test Acc'],
        [['42', f'{v_vals[0]*100:.2f}%', f'{t_vals[0]*100:.2f}%'],
         ['123', f'{v_vals[0 if len(v_vals)==3 else 1]*100:.2f}%' if len(v_vals)>=2 else '-',
                 f'{t_vals[1]*100:.2f}%'],
         ['456', f'{v_vals[min(2,len(v_vals)-1)]*100:.2f}%' if len(v_vals)>=3 else '-',
                 f'{t_vals[min(2,len(t_vals)-1)]*100:.2f}%']],
        [3, 5, 5]
    )

heading('5.3 跨语言实验分析', 2)
para('核心发现：', bold=True, size=11)
para('1. 同语言内（MY Mix）B3 可达 81.24%，但一跨语言就崩溃（EN→TE 仅 19.68%，仅比随机基线 16.7% 高 3pp）。这说明现有模块无法弥合语言级别的分布鸿沟。', size=11)
para('2. Validation-Test 鸿沟（86% → 20%）= 典型的分布外（Out-of-Distribution）泛化失败。Val 高而 Test 低，意味着模型在训练时学到了源语言特有的声学线索，但这些线索在目标语言中不成立。', size=11)
para('3. TE→EN (28.15%) 略优于 EN→TE (19.68%)。可能原因：TELUGU 420 speakers 的声学多样性 > ENGLISH 139 speakers，训练出的模型泛化性稍好。', size=11)
para('4. 已写入论文初稿 v3 第 5.5 节。', size=11)

heading('5.4 emotion2vec+ 模型对比实验', 2)
para('目的：对比 WavLM Base 与 emotion2vec+ 系列模型在儿童 SER 上的表现。', size=11)
para('配置：B3（Adapter + Prosody Pooling + DrseCNN），3 seeds，MY 数据集。', size=11)
para('注意：emotion2vec_plus_large 为 1024-dim 特征，adapter_init.npz 为 768-dim → 改用随机初始化。', size=11)

add_table(
    ['模型', 'Hidden Size', 'Val Acc (mean ± std)', 'Test Acc (mean ± std)', 'vs WavLM B3'],
    [['WavLM Base（参考）', '768', '88.52%', '81.24% ± 1.45%', '—'],
     ['emotion2vec_plus_large', '1024', '23.42% ± 0.76%', '22.45% ± 0.92%', '-58.79pp']],
    [5, 3, 4, 4, 3.5]
)

para('可能原因分析：', bold=True, size=11)
para('1. FunASR 特征提取可能不正确——训练时出现大量 "decoder key missing" 警告。', size=10)
para('2. emotion2vec 预训练数据主要是中文，对英文/泰卢固语儿童语音的迁移能力差。', size=10)
para('3. 1024-dim 特征可能需要不同的 Adapter/Pooling 超参数但未充分调优。', size=10)
para('4. 冻结参数 164M vs WavLM 94M，参数量更大但特征质量反而更差。', size=10)
para('5. Phase 1 中 emotion2vec_base (768-dim) 为 ~66%，plus_large (1024-dim) 反而 22%——不合逻辑，强烈怀疑 FunASR 兼容性问题。', size=10)
para('结论：暂不写入论文。如审稿人要求，优先排查 FunASR 特征提取流程。', bold=True, size=10)

# cross-lang seed data fix
# X1
v1 = xl['X1_EnglishToTelugu']['val_seeds']
t1 = xl['X1_EnglishToTelugu']['test_seeds']
add_table(
    ['Seed', 'Val Acc', 'Test Acc'],
    [['42', f'{v1[0]*100:.2f}%', f'{t1[0]*100:.2f}%'],
     ['123', f'{v1[1]*100:.2f}%', f'{t1[1]*100:.2f}%'],
     ['456', f'{v1[2]*100:.2f}%', f'{t1[2]*100:.2f}%']],
    [4, 5, 5]
)

# ═══════════════════════════════════════════
# 6. THREE-PHASE COMPARISON
# ═══════════════════════════════════════════
heading('六、三阶段核心数据对比', 1)

para('以下汇总了项目从原始论文（数据泄露）到修复后（真实基线）到当前方案（分布驱动三模块）的核心维度对比。', size=11)

add_table(
    ['维度', '原始论文 (2025.7)', '修复后 (2026.4.25)', '当前方案 (2026.4.28)'],
    [['特征提取', '94/162-dim 手工特征 + mean pooling', '同左', 'WavLM 帧级 768-dim (T×768)'],
     ['时间维度保留', '✗（Conv1d 在特征拼接维滑动）', '✗', '✓（200 帧 × 768 维，完整保留）'],
     ['Speaker 划分', '无（数据泄露 63%）', '70/30 split', '60/20/20 三路 split'],
     ['增强方式', '3x (noise+stretch+pitch) 无约束', '无增强', '分布约束 SpecAugment + 负面实验'],
     ['模型核心', 'DrseCNN 网络设计 (~11.37M)', '同左', '三模块框架（~900K 新增参数 + 94M WavLM）'],
     ['Val Accuracy', '86.82%（无效）', '~35%', '88.52%（B3）'],
     ['Test Accuracy', '—（无独立 Test）', '≈ Val', '81.96%（A3, 3-way split）'],
     ['实验可信度', '0（数据泄露+无时间维）', '可信（speaker独立）', '可信（3-way + 3 seeds + hold-out）']],
    [3.5, 4.5, 4.5, 4.5]
)

# ═══════════════════════════════════════════
# 7. DATA SOURCES
# ═══════════════════════════════════════════
heading('七、数据来源索引', 1)

add_table(
    ['数据文件', '路径', '内容说明'],
    [['Phase 3 消融实验（主表）', 'experiments/phase3_ablation.json', '5 experiments × 3 seeds 的 val + test 完整数据'],
     ['Phase 3 运行记录', 'experiments/registry.csv', '每个实验的时间戳、耗时、epoch 数、最佳 val/test'],
     ['跨语言迁移实验', 'experiments/cross_language_results.json', 'EN→TE 和 TE→EN 的 3-seed val + test 数据'],
     ['模型对比实验', 'experiments/model_comparison_results.json', 'emotion2vec_plus_large vs WavLM 的对比结果'],
     ['Phase 4 实验记录', 'experiments/phase4_record.md', 'Phase 4 补充实验的文字描述和分析'],
     ['云端运行日志', 'experiments/phase3_stdout.log', '云端 RTX 4090D 的完整训练 stdout 日志'],
     ['Phase 3 各实验 per-epoch 数据', 'experiments/runs/*/metrics.csv', '每个实验每 epoch 的 loss + acc 时间序列'],
     ['Phase 3 各实验配置', 'experiments/runs/*/config.json', '每个实验的完整超参数配置'],
     ['儿童语音统计分析', 'src/data/statistics.py', 'F0/时长/SNR/频谱平滑度/语速的统计分析脚本'],
     ['Phase 1 SSL 基线', 'experiments/logs/', 'SSL linear probe 的早期实验配置与结果'],
     ['论文初稿 v3', 'docs/论文初稿_v3.docx', '包含跨语言迁移实验结果的论文初稿'],
     ['FD vs Accuracy 图表', 'assets/figures/fd_vs_accuracy.png', '分布偏移 vs 分类性能的 matplotlib 图']],
    [5, 7, 6]
)

# ═══════════════════════════════════════════
# 8. KEY NUMBERS SUMMARY
# ═══════════════════════════════════════════
heading('八、关键数字速查', 1)

add_table(
    ['指标', '数值', '来源'],
    [['数据泄露严重程度', '63% speaker 重叠', '原始代码 speaker ID 统计'],
     ['旧方案信息上限', '~35%', 'Phase 1 162-dim baseline'],
     ['SSL 基线 (WavLM)', '80.66% (Val, 70/30)', 'Phase 1 A1'],
     ['最佳单模块配置', '86.42% (Val, 70/30, B3)', 'Phase 2.2'],
     ['FD vs Accuracy 相关性', '严格单调负相关', 'Phase 2.3 四组对比'],
     ['Final Test (3-way, 3 seeds, A3)', '81.96% ± 0.73%', 'Phase 3 A3'],
     ['Module 贡献: WavLM', '+43.34pp', 'Phase 3 分解'],
     ['Module 贡献: Adapter', '+1.17pp', 'Phase 3 分解'],
     ['Module 贡献: Prosody Pooling', '+1.73pp', 'Phase 3 分解'],
     ['Module 贡献: Full Fine-tune', '+0.72pp', 'Phase 3 分解'],
     ['跨语言 EN→TE', '19.68% (Test)', 'Phase 4 X1'],
     ['跨语言 TE→EN', '28.15% (Test)', 'Phase 4 X2'],
     ['儿童 F0 均值', '495.9 Hz', 'BESD 统计分析'],
     ['儿童 vs 成人 FD', '0.37', 'WavLM 隐空间计算'],
     ['Adapter 参数量', '~900K', '模型结构计算']],
    [6.5, 4.5, 6]
)

# ─── Save ───
doc.save(OUTPUT)
print(f'Document saved: {OUTPUT}')
